"""Generate 概要设计书.docx for student project."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

for _ in range(6):
    doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ODPlatform 工业安全帽佩戴检测\n概要设计说明书')
run.font.size = Pt(22)
run.font.bold = True

doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'版本: V1.0\n日期: {datetime.date.today()}\n组别: 第5组').font.size = Pt(12)

doc.add_page_break()

sections = [
    ("1. 项目概述", [
        ("1.1 项目背景", "本项目旨在构建一个与数据集解耦、与业务场景解耦的通用目标检测工程平台——ODPlatform。V1.0以工业安全帽佩戴检测为落地场景，采用YOLOv8算法，实现从原始数据到端侧推理的全链路覆盖。"),
        ("1.2 项目目标", "端到端覆盖流程: 原始数据->格式转换->数据验证->配置管理->训练->评估->推理->可视化。平台不绑定特定数据集，通过配置文件驱动。推理引擎通过OutputSink接缝支持桌面GUI、Web界面等多种前端。"),
        ("1.3 数据集", "使用公开安全帽佩戴检测数据集(SHWD风格)，包含7581张标注图片，2个类别(hat/person)，原始格式为PASCAL VOC。"),
    ]),
    ("2. 技术架构", [
        ("2.1 技术栈", "语言: Python 3.10+ | 深度学习: PyTorch + Ultralytics (YOLOv8) | 图像处理: OpenCV + Pillow | 配置校验: Pydantic v2 | 桌面GUI: tkinter | Web界面: FastAPI + WebSocket + HTML5/CSS3/JS | 测试: pytest"),
        ("2.2 架构约束", "1) 业务模块只向下依赖common，不反向、不横向耦合 (CON-06)\n2) 路径统一走paths.py，无硬编码 (CON-05)\n3) 推理包零宿主依赖 (NFR-EXT-03)\n4) 对外API通过__all__声明 (NFR-ENG-04)"),
        ("2.3 分层架构", "展示层: CLI / desktop(tkinter) / web(FastAPI)\n业务层: training / evaluation / inference / data_pipeline / data_validation / runtime_config / frame_source / visualization\n基础层: common (paths / logging / utils / constants)"),
    ]),
    ("3. 模块详细设计", [
        ("3.1 common 基础设施层 (高晨曦)", "paths.py: 工作区根路径定位(通过.odp-workspace标记)，所有目录路径统一管理。logging_utils.py: 按级别着色日志(console+file)，支持CJK。result.py: 不可变评估指标数据类。model_utils.py: 模型路径解析。constants.py: 全局常量。"),
        ("3.2 data_pipeline 数据流水线 (高文焱)", "convert/: 格式转换器注册(支持PASCAL VOC/COCO/YOLO互转)。split/: 随机/分层切分策略，materializer物化。orchestrator.py: 编排器串联Convert->Split->Materialize->Write YAML全流程。"),
        ("3.3 data_validation 数据验证 (高文焱)", "四项检查: yaml_schema(配置格式)、pair_existence(图像标注配对)、label_format(标注格式)、split_uniqueness(切分无重叠)。一次扫盘构建snapshot，聚合执行，生成可读报告。"),
        ("3.4 runtime_config 运行配置 (刘睿煊)", "基于Pydantic v2的训练/验证/推理配置模型。ConfigGenerator反射式YAML模板生成器。字段级校验(ge/gt/范围)，支持CLI参数覆盖。"),
        ("3.5 training 训练编排 (刘睿煊)", "TrainService: 配置日志->数据校验->底层训练->权重归档->审计快照。支持CLI覆盖、学术绘图风格、自动设备检测。实测: RTX 4090, 100 epoch, mAP50=0.939。"),
        ("3.6 evaluation 评估 (陈瀚鹏)", "ValService: 解析权重->解析数据集->YOLO val()->收集指标。输出mAP50/mAP50-95/Precision/Recall。实测: mAP50=0.936, Precision=0.937, Recall=0.895。"),
        ("3.7 inference 推理流水线 (陈瀚鹏)", "三线程流水线: 采集(FrameSource)->推理(YOLO)->输出(OutputSink)。Queue解耦，CancelToken优雅退出。OutputSink抽象接缝不依赖任何前端框架。"),
        ("3.8 frame_source 统一帧源 (陈瀚鹏)", "create_frame_source()自动识别: 摄像头序号/RTSP/图片文件/图片目录/视频文件。支持ImageSource/VideoSource/CameraSource统一迭代接口。"),
        ("3.9 visualization 可视化 (王智鹏)", "BeautifyVisualizer: 边界框+中文标签+置信度。Pillow渲染CJK，OpenCV回退英文。可配置调色板，零宿主依赖。"),
        ("3.10 desktop 桌面GUI (王智鹏)", "tkinter实现(Windows原生)，深色主题。支持摄像头/图片/视频三种源。TkDisplaySink桥接InferService，线程安全主线程帧更新。"),
        ("3.11 web-backend Web界面 (王智鹏)", "FastAPI+WebSocket单文件应用，内嵌HTML/CSS/JS。三大模式: 实时摄像头(WebSocket逐帧)、图片上传(拖拽)、视频文件。深色美学主题，实时统计面板。"),
    ]),
    ("4. 数据流设计", [
        ("4.1 训练数据流", "原始数据集(PASCAL VOC) -> odp-transform(转换+切分) -> data/{train,val,test}/ -> odp-validate(四维质检) -> data.yaml -> odp-train(YOLOv8) -> best.pt"),
        ("4.2 推理数据流", "输入源(摄像头/图片/视频/RTSP) -> FrameSource(统一帧迭代) -> YOLO推理 -> BeautifyVisualizer(画框+标签) -> OutputSink(显示/保存/Web推送)"),
        ("4.3 Web实时推理流", "浏览器摄像头/视频 -> canvas.toBlob(JPEG) -> WebSocket二进制 -> cv2解码 -> YOLO推理 -> 画框 -> base64 -> WebSocket JSON -> 浏览器展示"),
    ]),
    ("5. CLI接口定义", [
        ("5.1 命令列表", "odp-init: 初始化运行时目录\nodp-reset: 清理运行时产物\nodp-transform: 数据转换+切分 (--dataset --format --classes)\nodp-validate: 数据质检 (--dataset --split)\nodp-gen-config: 生成配置模板\nodp-train: 训练 (--model --data --epochs --batch)\nodp-eval: 评估 (--weights --data)\nodp-infer: 推理 (--source --weights)"),
    ]),
    ("6. 测试策略与结果", [
        ("6.1 测试概览", "单元测试12项 + 集成测试4项 + 端到端测试2项 = 共18项，全部PASS。"),
        ("6.2 关键指标", "数据验证: ERROR=0 WARNING=0 INFO=4\n训练: mAP50=0.939, mAP50-95=0.616\n评估: mAP50=0.936, Precision=0.937, Recall=0.895\n推理: 单帧约10ms(RTX 3070), Web实时FPS约15-20"),
    ]),
    ("7. 项目交付物清单", [
        ("7.1 交付物", "1) 项目源码 (apps/platform/ + apps/desktop/ + apps/web-backend/)\n2) 训练模型 (best.pt, mAP50=0.939)\n3) 项目运行录屏 (Web界面+桌面GUI)\n4) 概要设计书 (本文档)\n5) 结合测试报告\n6) 人员分工表\n7) 答辩PPT\n8) 团队贡献度表"),
    ]),
]

for sec_title, subs in sections:
    h = doc.add_heading(sec_title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(26, 82, 118)
    for sub_title, content in subs:
        doc.add_heading(sub_title, level=2)
        for line in content.split('\n'):
            doc.add_paragraph(line.strip())

base = r'G:\CLAUDE CODE_PROJECT\ODP\【学生成果物】第5组_工业安全帽佩戴检测-高晨曦_王智鹏_刘睿煊_陈瀚鹏_高文焱\03.项目文件'
doc.save(base + '\\概要设计书_第5组_工业安全帽佩戴检测.docx')
print('概要设计书 done')
